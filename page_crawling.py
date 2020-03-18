from selenium import webdriver
import docx
from bs4 import BeautifulSoup
import time

class Applicant:
    def __init__(self):
        self.name = ""
        self.info = ""
        self.part = ""
        self.q1 = ""
        self.q2 = ""
        self.q3 = ""
        self.q4 = ""
        self.date = ""

    def todocx(self):
        doc = docx.Document()

        str_q1 = str(len(self.q1))
        str_q2 = str(len(self.q2))
        str_q3 = str(len(self.q3))
        str_q4 = str(len(self.q4))

        para = doc.add_paragraph()
        run = para.add_run(self.info)
        run.font.name = 'Arial'
        run.bold = True

        para = doc.add_paragraph()
        run = para.add_run("1. 지원 분야를 선택해 주세요. (하나의 분야에만 지원 가능)")
        run.font.name = 'Arial'
        run.bold = True

        para = doc.add_paragraph()
        run = para.add_run(self.part)
        run.font.name = 'Arial'

        para = doc.add_paragraph()
        run = para.add_run("2. 자기 소개와 함께 멋쟁이사자처럼에 지원하게 된 동기를 적어 주세요. [글자수] "+str_q1+" 자")
        run.font.name = 'Arial'
        run.bold = True

        para = doc.add_paragraph()
        run = para.add_run(self.q1)
        run.font.name = 'Arial'

        para = doc.add_paragraph()
        run = para.add_run("3. 기억에 남는 프로그래밍 경험과 느낀 점에 대해 적어 주세요. 만약 프로그래밍 경험이 없다면, 어떤 언어를 배우고 싶은지 적어 주세요. [글자수] "+str_q2+" 자")
        run.font.name = 'Arial'
        run.bold = True

        para = doc.add_paragraph()
        run = para.add_run(self.q2)
        run.font.name = 'Arial'

        para = doc.add_paragraph()
        run = para.add_run("4. 멋쟁이사자처럼에서 이루고 싶은 목표 및 만들고 싶은 서비스에 대해 설명해 주세요. [글자수] "+str_q3+" 자")
        run.font.name = 'Arial'
        run.bold = True

        para = doc.add_paragraph()
        run = para.add_run(self.q3)
        run.font.name = 'Arial'

        para = doc.add_paragraph()
        run = para.add_run("5. 참여했던 팀 활동 중 가장 기억에 남는 경험과 느낀 점을 적어 주세요. [글자수] "+str_q4+" 자")
        run.font.name = 'Arial'
        run.bold = True

        para = doc.add_paragraph()
        run = para.add_run(self.q4)
        run.font.name = 'Arial'

        para = doc.add_paragraph()
        run = para.add_run("7. 면접이 가능하신 날짜를 체크해 주세요.")
        run.font.name = 'Arial'
        run.bold = True

        para = doc.add_paragraph()
        run = para.add_run(self.date)
        run.font.name = 'Arial'

        fileName = self.name + ".docx"
        doc.save(fileName)

driver = webdriver.Chrome("./chromedriver")
driver.get("http://www.likelion-mju.com/admin/login/?next=/admin/")

# id
id_box = driver.find_element_by_css_selector('input#id_username')
id_box.send_keys("iamadmin")

# password
pass_box = driver.find_element_by_css_selector('input#id_password')
pass_box.send_keys("iamadmin!")

driver.find_element_by_css_selector('div.submit-row input').click()

time.sleep(2)

driver.get("http://www.likelion-mju.com/apply/list/submit")
soup = BeautifulSoup(driver.page_source, 'html.parser')

applicants = soup.select("ul.list-group li")
print(len(applicants))
print(applicants)

all = []
# tmp = driver.find_elements_by_css_selector('li.list-group-item a')
for i in range(0, len(applicants)):
    driver.find_elements_by_css_selector('li.list-group-item a')[i].click()
    # driver.find_element_by_css_selector("ul.list-group a:nth-of-type("+str(i)+")").click()
    time.sleep(2)
    html = BeautifulSoup(driver.page_source, 'html.parser')
    ap = Applicant()
    info = html.select('div.user-info')
    ap.name = info[0].text[-8:]+'_'+info[1].text[-3:]
    print(ap.name)
    for i in info:
        ap.info = ap.info + i.text + "\n"
    ap.part = html.select_one('input#field[checked]').attrs['value']
    ap.q1 = html.select_one('textarea#answer1').text + "\n"
    ap.q2 = html.select_one('textarea#answer2').text + "\n"
    ap.q3 = html.select_one('textarea#answer3').text + "\n"
    ap.q4 = html.select_one('textarea#answer4').text + "\n"
    ap.date = html.select_one('input#date[checked]').attrs['value']
    all.append(ap)
    driver.get("http://www.likelion-mju.com/apply/list/submit")

for a in all:
    a.todocx();

print(all)
print(ap.date)
print(ap)
driver.close()